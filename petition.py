from openai import OpenAI
from docx import Document
from docx.shared import Inches
import gspread

#Access Spreadsheet with responses
gc = gspread.service_account()

sh = gc.open("ATG Intake Try 1 (Responses)")


#Ask for client number which is the row that the answers are.
client_num= input ("Please enter the client number:")


#VARIABLE LIST
#These will need to be added by a person for now
county = "Washington"
court_file_num = "3XXXX3"
hearing_date = "When will we know?"

#These are added by the spreadsheat or programmtically added from spreadsheet data

client_first_name = sh.sheet1.acell(str("B" + client_num )).value
client_middle_name = sh.sheet1.acell(str("C" + client_num )).value
client_last_name = sh.sheet1.acell(str("D" + client_num )).value
client_full_name = client_first_name + " " + client_middle_name + " " + client_last_name


#TODO?: Split this into a proper list?
other_legal_names = sh.sheet1.acell(str("E" + client_num )).value

client_birth_date = sh.sheet1.acell(str("F" + client_num )).value
client_street_address = sh.sheet1.acell(str("G" + client_num )).value
client_city_state_zip = sh.sheet1.acell(str("H" + client_num )).value
client_other_addresses = sh.sheet1.acell(str("I" + client_num )).value

client_reasons_for_expunge = sh.sheet1.acell(str("J" + client_num )).value
client_steps_taken_4rehab = sh.sheet1.acell(str("K" + client_num )).value
client_things_2help = sh.sheet1.acell(str("L" + client_num )).value

prompt_question = "Compose a paragraph answering this question in first person: "
llm_role = "You are an average person seeking an expungment from a judge."


print(client_first_name)


#Functions=
def make_bold(para):
    run = para.runs[0]
    run.bold = True
    return para

def make_italic(para):
    run = para.runs[0]
    run.italic = True
    return para


#Start the document
document = Document()

#document.add_heading('Document Title', 0)

p1 = document.add_paragraph('')
p1.add_run('State of Minnesota').bold = True
p1.add_run('                              ')
p1.add_run('                              ')
p1.add_run('                              ')
p1.add_run('District Court').bold = True


p2 = document.add_paragraph('')
p2.add_run('County of: ').bold = True
#ADD in County programmitically
p2.add_run(county)

p3 = document.add_paragraph('Court File Number: ')
make_bold(p3)
#ADD in Court FIle Number

p4 = document.add_paragraph('Judicial District: ')
make_bold(p4)

p5 = document.add_paragraph('Case Type: ')
make_bold(p5)
p5.add_run("Criminal")

p6 = document.add_paragraph('State of Minnesota ')

p7 = document.add_paragraph('Plaintiff')
make_italic(p7)

p8 = document.add_paragraph(' VS ')
make_bold(p8)

p9 = document.add_paragraph('')
p9.add_run(client_full_name)

p10 = document.add_paragraph('Defendant')
make_italic(p10)


document.add_heading('Notice of Hearing and Petition for Expungement (EXP102)', level=1)
document.add_heading('Minn. Stat. § 609A.03 or Inherent Authority', level=3)
#document.add_paragraph('Minn. Stat. § 609A.03 or Inherent Authority', style='Intense Quote')

p11 = document.add_paragraph("Notice to Law Enforcement / Government Agency / Prosecutor: ")
make_bold(p11)
p11.add_run ('Any objection to an expungement in this case shall be filed with the court as soon as possible, and within 60 days').bold = True

document.add_heading('Hearing Information', level=2)

p12 = document.add_paragraph('The Hearing in this case is scheduled on:')

p13 = document.add_paragraph(hearing_date)

p14 = document.add_paragraph('This hearing will be:')

p15 = document.add_paragraph('REMOTE USING ZOOM')

p16 = document.add_paragraph('Zoom Meeting ID:')

p17 = document.add_paragraph('Passcode:')

document.add_page_break()


#HERE is where the petition for expungmenet begins


document.add_heading('Petition for Expungement', level=2)


pe1 = document.add_paragraph('1. List your full name:')
make_bold(pe1)


pe1a= document.add_paragraph('  First: ')
make_bold(pe1a)
pe1a.add_run(client_first_name)

pe1b= document.add_paragraph('  Middle: ')
make_bold(pe1b)

pe1b.add_run(client_middle_name)

pe1c= document.add_paragraph('  Last: ')
make_bold(pe1c)

pe1c.add_run(client_last_name)


pe2 = document.add_paragraph('2. List any other legal names or aliases you have been known as: ')
make_bold(pe2)

## TODO: ADD TEXT IF THE ANSWER IS NONE

pe2.add_run(other_legal_names)


pe3 = document.add_paragraph('3. List your date of birth: ')
make_bold(pe3)
pe3.add_run(client_birth_date)


pe4 = document.add_paragraph('4. List your current address: ')
make_bold(pe4)
pe4a = document.add_paragraph('  Street Address: ')
make_bold(pe4a)
pe4a.add_run(client_street_address)
pe4b = document.add_paragraph('  City, State, Zip: ')
make_bold(pe4b)
pe4b.add_run(client_city_state_zip)


pe5 = document.add_paragraph('5. List all the other addresses you have lived at since the date of the offence for which you are seeking an expungement: ')
make_bold(pe5)
pe5a = document.add_paragraph('  Check this box if you have only lived at your current address (listed in #4) since the date of the offense')
make_bold(pe5a)
#TODO: ADD box here


pe5b = document.add_paragraph('  See attached affidavit for answer')

pe6 = document.add_paragraph('6. Reasons for this Request: ')
make_bold(pe6)
pe6a = document.add_paragraph('  I am asking for an expungement because:')
make_bold(pe6a)
pe6b = document.add_paragraph('  See attached affidavit for answer')


pe7 = document.add_paragraph('7. Criminal Record.')
make_bold(pe7)
pe7a = document.add_paragraph('  THIS IS WHERE THE TABLE WILL GO TO BE ADDED BY THE LAWYER')


pe8 = document.add_paragraph('8. Past Requests. Have you ever asked for an expungemenet, pardon, or sealing of a criminal record before?')
make_bold(pe8)
pe8a = document.add_paragraph('  CLIENT ANSWER HERE YES OR NO')
#TODO: MUST ADD THIS TO THE FORM


pe8b = document.add_paragraph('  If yes, list each other request for an expungement, pardon, or sealing of a criminal record you have made:')
make_bold(pe8b)
pe8b.add_run(client_other_addresses)

pe9 = document.add_paragraph('9. Qualification for Expungement.')
make_bold(pe9)
pe9a = document.add_paragraph('  Why do you qualify for an expungement?')
make_bold(pe9a)

pe10 = document.add_paragraph('10. Offense Details. What is the offense you want to expunge?')
make_bold(pe10)
pe10a = document.add_paragraph('  Case#:')
pe10b = document.add_paragraph('  Jurisdiction/City where the offense occured:')
pe10c = document.add_paragraph('  Type of offense:')
pe10d = document.add_paragraph('  Date of offense:')

pe11 = document.add_paragraph('11. Victims. Were there any identifiable victims in this case?')
make_bold(pe11)
#Add Yes or No here
pe11a = document.add_paragraph('  If Yes, list the names of the victims: ')
make_bold(pe11a)

pe12 = document.add_paragraph('12.	Protection, Restraining, or No-Contact Orders. Is there now, or has there ever been, an Order for Protection, Restraining Order, or other No-Contact Order prohibiting you from contacting the victims? #:')
make_bold(pe12)
pe12a = document.add_paragraph('  YES OR NO OR N/A BASED ON CLIENT ANSWERS')

pe13 = document.add_paragraph('13.	Personal Rehabilitation.  Describe what steps you have taken since the time of the offense toward personal rehabilitation, including treatment, work, or other personal history that demonstrates rehabilitation.   ')
make_bold(pe13)
pe13a = document.add_paragraph('  I have taken the following steps toward personal rehabilitation:')
make_bold(pe13a)
pe13b = document.add_paragraph('  See attached affidavit for answer')


pe14 = document.add_paragraph('14. Do you want to ask the court to seal any private or confidential data submitted by the responding agencies or other jurisdictions? [Minn. Stat. § 609A.03, subd. 3(d)] ')
make_bold(pe14)
pe14a = document.add_paragraph('  Yes')


pe15 = document.add_paragraph(' 15.	Do you want to ask that, if an expungement is ordered, each agency and jurisdiction that receives the order must send a letter to you at the address provided confirming receipt of the expungement order and that the record has been expunged? [Minn. Stat. § 609A.03, subd. 8(b)]: ')
make_bold(pe15)
pe15 = document.add_paragraph('  Yes')

pe16 = document.add_paragraph('  16.	Mitigating or Aggravating Factors. Explain any mitigating or aggravating factors relating to the underlying crime, including your level of participation, the context and circumstances of the underlying crime, and what risk, if any, you pose to individuals or society')
make_bold(pe16)

pe16 = document.add_paragraph('  See attached affidavit for answer')


document.add_page_break()
pe17 = document.add_paragraph('I declare under penalty of perjury that everything I have stated in this document is true and correct.  Minn. Stat. § 358.116 ')
make_bold(pe17)

#to be filled in by the lawyer
pe18 = document.add_paragraph('Date:')
make_bold(pe18)
pe19 = document.add_paragraph('Signature: ')
make_bold(pe19)
pe20 = document.add_paragraph('County and State where signed:')
make_bold(pe20)
pe21 = document.add_paragraph('Name:')
make_bold(pe21)
pe22 = document.add_paragraph('Address:')
make_bold(pe22)
pe23 = document.add_paragraph('City/State/Zip:')
make_bold(pe23)
pe24 = document.add_paragraph('Phone:')
make_bold(pe24)
pe25 = document.add_paragraph('Email:')
make_bold(pe25)


document.save('petitionAI.docx')

#Here we generate the affadavit
##HERE IS WHERE THE Affadvait is


affa = Document()

client = OpenAI()

completion = client.chat.completions.create(
  model="gpt-4",
  messages=[
    {"role": "system", "content": llm_role },
    {"role": "user", "content": str(prompt_question + "What are the reasons I am looking for an expungement? Your reasons are: " + client_reasons_for_expunge ) }
  ],
  temperature = 0.1
)


print(completion.choices[0].message.content)
print("  ")



q6_answer = completion.choices[0].message.content


completion2 = client.chat.completions.create(
  model="gpt-4",
  messages=[
    {"role": "system", "content": llm_role},
    {"role": "user", "content": str( prompt_question + "Describe what steps you have taken since your criminal offense for personal rehabilitation? Your steps taken are: " + client_steps_taken_4rehab ) }
  ],
  temperature = 0.1
)


print(completion2.choices[0].message.content)
print("  ")

q13_answer = completion.choices[0].message.content


completion3 = client.chat.completions.create(
  model="gpt-4",
  messages=[
    {"role": "system", "content": llm_role},
    {"role": "user", "content": str (prompt_question + "What factors should grant you an expungement?' Your relevant factors are: " + llm_role)}
  ],
    temperature = 0.0
)


print(completion3.choices[0].message.content)
print("  ")

q16_answer = completion3.choices[0].message.content



a1 = affa.add_paragraph('')
a1.add_run('State of Minnesota').bold = True
a1.add_run('                              ')
a1.add_run('                              ')
a1.add_run('                              ')
a1.add_run('District Court').bold = True

a2 = affa.add_paragraph('')

a2.add_run('County of: ').bold = True
#ADD in County programmitically
a2.add_run(county)

a3 = affa.add_paragraph('Court File Number: ')
make_bold(a3)
#ADD in Court FIle Number

a4 = affa.add_paragraph('Judicial District:  ')
make_bold(a4)

a5 = affa.add_paragraph('Case Type: Criminal')

a6 = affa.add_paragraph('State of Minnesota ')
make_bold(a6)

a7 = affa.add_paragraph('Plaintiff')
make_italic(a7)


a8 = affa.add_paragraph(' VS ')
make_bold(a8)

a9 = affa.add_paragraph('')
a9.add_run(client_full_name)

a10 = affa.add_paragraph('Defendant')
make_italic(a10)


a11 = affa.add_paragraph('I, ' + client_full_name + ", state under the penalty of prejury that the following stateements are true and correct to the bset of my knowledge, belief, and recollection: ")

a12 = affa.add_paragraph('This affidavit is to serve as further explanation of the case I wish to expunge and demonstrate the progress and rehabilitative efforts I have made since then.')

a13 = affa.add_paragraph('Answer to question 5')
make_bold(a13)
a13a = affa.add_paragraph('TO BE COMPLETED')


a14 = affa.add_paragraph('Answer to question 6')
make_bold(a14)

a15 = affa.add_paragraph(q6_answer)

a16 = affa.add_paragraph('Answer to question 13')
make_bold(a16)

a17 = affa.add_paragraph(q13_answer)

a18 = affa.add_paragraph('Answer to question 16')
make_bold(a18)

a19 = affa.add_paragraph(q16_answer)


a20 = affa.add_paragraph('I declare under penalty of perjury that everything I have stated in this document is true and correct')

make_bold(a20)


a21 = affa.add_paragraph('Dated: ')
make_bold(a21)

a22= affa.add_paragraph('Signature: /s/')

make_bold(a22)

a23 = affa.add_paragraph('                         ' + client_full_name)




affa.save('affadavitAI.docx')

print ("Complete")

#Relevant Variables
#client_reasons_for_expunge 
#client_steps_taken_4rehab
#client_things_2help 





